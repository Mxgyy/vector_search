import streamlit as st
import os
import subprocess
import shutil
from tkinter import Tk
from tkinter.filedialog import askopenfilename, askopenfilenames
import tkinter as tk
from tkinter import filedialog
from embedding import Myembedding_funcxxx
import shutil
import os
from langchain.schema import Document
from langchain.vectorstores import FAISS
from langchain_text_splitters import RecursiveCharacterTextSplitter
import docx
import random

# 设置 tkinter
root = tk.Tk()
root.withdraw()
# 使文件选择对话框显示在其他窗口之上
root.wm_attributes('-topmost', 1)

# # 确保documents文件夹存在
# if not os.path.exists('documents'):
#     os.makedirs('documents')

if 'query' not in st.session_state:
    st.session_state['query'] = ""


# 设置页面标题
st.set_page_config(page_title="智能文档内容检索")

# st.markdown(
#     """
#     <style>
#     button[kind="primary"] {
#         width: 100%;
#     }

#     .st-key-res_button0 > div:first-child > button:first-child{
#         width: 100%;
#         color: blue;
#         background: blue;
#     }
#     .st-key-res_button1 > div:first-child > button:first-child{
#         width: 100%;
#         color: blue;
#         background: blue;
#     }

#     </style>
#     """,
#     unsafe_allow_html=True,
# )

st.markdown(
    """
    <style>
    button[kind="primary"] {
        width: 100%;
    }

    </style>
    """,
    unsafe_allow_html=True,
)

# st.markdown("<style>.button1 { /* Button 1 CSS attributes*/ }</style>", unsafe_allow_html=True)
# button1_clicked = st.button("Button 1")
# st.markdown("<style>.button2 { /* Button 2 CSS attributes*/ }</style>", unsafe_allow_html=True)
# button2_clicked = st.button("Button 2")

document_path = ".\\resources\\documents"
file_list_path = ".\\resources\\file_list.txt"
vdb_path = ".\\resources\\VDB.db"


def write_css(key):
    st.markdown(
        f"""
        <style>
        .st-key-{key}{{
            width: 100%;
            white-space: nowrap;
            overflow: hidden;
            text-overflow: ellipsis;
            text-align:left;
            font-size: 13px;
            padding: 0px;
            margin: 0px;
            line-height:1;
        }}
        .st-key-{key} > div:first-child{{
            width: 100%;
            white-space: nowrap;
            overflow: hidden;
            text-overflow: ellipsis;
            text-align:left;
            font-size: 13px;
            padding: 0px;
            margin: 0px;
            line-height:1;
        }}
        .st-key-{key} > div:first-child > button:first-child{{
            width: 100%;
            white-space: nowrap;
            overflow: hidden;
            text-overflow: ellipsis;
            text-align:left;
            font-size: 13px;
            padding: 0px;
            margin: 0px;
            line-height:1;
        }}
        .st-key-{key} > div:first-child > button:first-child >div:first-child{{
            width: 100%;
            white-space: nowrap;
            overflow: hidden;
            text-overflow: ellipsis;
            text-align:left;
            font-size: 13px;
            padding: 0px;
            margin: 0px;
            line-height:1;
        }}
        </style>
        """,
        unsafe_allow_html=True,
    )

def copy_file_to_directory(source_file_path, target_directory):
    # Check if the source file exists
    if not os.path.isfile(source_file_path):
        raise FileNotFoundError(f"The source file {source_file_path} does not exist.")
    # Check if the target directory exists, if not, create it
    if not os.path.isdir(target_directory):
        os.makedirs(target_directory)
    # Construct the target file path
    target_file_path = os.path.join(target_directory, os.path.basename(source_file_path))
    # Copy the file
    shutil.copy2(source_file_path, target_file_path)
    print(f"File copied successfully to {target_file_path}")

def read_filenames_to_list(file_path):
    """
    读取一个文件，其中每行是一个文件名，然后返回这些文件名的列表。
    
    :param file_path: 包含文件名的文件路径。
    :return: 文件名列表。
    """
    with open(file_path, 'r') as file:
        filenames = [line.strip() for line in file]
    return filenames

def find_unique(list1, list2):
    """
    找出第一个列表有而第二个列表没有的元素。

    :param list1: 第一个列表。
    :param list2: 第二个列表。
    :return: 一个列表，包含第二个列表中有而第一个列表中没有的元素。
    """
    # 将两个列表转换为集合
    set1 = set(list1)
    set2 = set(list2)
    
    # 使用集合的差集操作找出第二个列表独有的元素
    unique_in_list = set1 - set2
    
    # 将结果转换回列表
    return list(unique_in_list)

def split_docx(filepath, chunk_size=256, chunk_overlap=0, separators=[".","。","\n","\n\n","?","!"]):
    doc = docx.Document(filepath)
    raw_document = "\n".join([para.text for para in doc.paragraphs])
    splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap,separators=separators)
    return splitter.split_text(raw_document)

def write_list_to_file(lst, file_name):
    with open(file_name, 'w') as file:
        for item in lst:
            file.write(str(item) + '\n')

if __name__ == "__main__":
    # 初始化
    if not os.path.exists(document_path):
        os.makedirs(document_path)
    if not os.path.exists(file_list_path):
        with open(file_list_path, "w") as f:
            f.write("")
    #加载FAISSDB
    vector_store = None
    if not os.path.exists(vdb_path):
        vector_store = FAISS.from_documents([Document(page_content="hello-world", metadata={"file": "test"}),Document(page_content="world-hello", metadata={"file": "test"})], Myembedding_funcxxx())
    else:
        vector_store = FAISS.load_local(vdb_path, Myembedding_funcxxx(),allow_dangerous_deserialization=True)


    # 主页面内容
    st.title("智能文档内容检索系统")
    # 侧边栏
    with st.sidebar:
        st.title("文档仓库管理")
        # 创建两列
        col1, col2 = st.columns(2)
        # 桌面文件夹按钮
        with col1:
            if st.button("添加文档"):
                file_paths = filedialog.askopenfilenames(master=root)
                if file_paths:
                    for file_path in file_paths:
                        copy_file_to_directory(file_path, document_path)
        with col2:
            if st.button("删除文档"):
                file_paths = filedialog.askopenfilenames(master=root,initialdir=document_path)
                if file_paths:
                    for file_path in file_paths:
                        try:
                            os.remove(file_path)
                        except Exception as e:
                            st.error(f"文件删除失败: {str(e)}")
        
        if st.button(" 构  建  索  引 ",type="primary"):
            # 在这里添加构建索引的代码
            pre_list = read_filenames_to_list(file_list_path)
            current_list = os.listdir(document_path)

            files2add = find_unique(current_list, pre_list)
            files2del = find_unique(pre_list, current_list)

            progress_bar = st.sidebar.progress(0)
            #删除文件索引
            for file in files2del:
                while True:
                    para2delete = vector_store.search("*", k=20,search_type="similarity",filter={"file": file})
                    if len(para2delete) == 0:
                        break
                    id2delete = [para.id for para in para2delete]
                    vector_store.delete(id2delete)
                    vector_store.save_local(vdb_path)
            #添加文件索引
            for file in files2add:
                if not file.endswith(".docx"):
                    continue
                texts = split_docx(f'{document_path}\\{file}')
                documents = []
                for text in texts:
                    doc = Document(page_content=text, metadata={"file": file})
                    documents.append(doc)
                if len(documents) > 0:
                    vector_store.add_documents(documents)
                    vector_store.save_local(vdb_path)
                
            #保存记录
            write_list_to_file(current_list,file_list_path)

            progress_bar.progress(100)
            st.write("索引构建完成！")
        
        st.title("检索结果显示")

    query = st.chat_input("请输入希望搜索的语句：")
    if query != "" and query != st.session_state['query'] and query != None:
        print(query)
        st.session_state['query'] = query
        res = vector_store.search(query, k=20,search_type="similarity")
        #for doc in res:
        for i in range(len(res)):
            doc = res[i]
            filename = doc.metadata['file']
            content = doc.page_content
            write_css(f"res_button{i}")
            if st.sidebar.button(f'{i}.{filename}',key=f"res_button{i}"):
                text_input = st.text_area(filename, height=350, value=content)
                if st.button("打开文件"):
                    subprocess.run(['start', f'{document_path}\\{filename}'], shell=True)
    
    elif query == None and st.session_state['query'] != "":
        query = st.session_state['query']
        st.session_state['query'] = query
        res = vector_store.search(query, k=20,search_type="similarity")
        for doc in res:
            filename = doc.metadata['file']
            content = doc.page_content
            if st.sidebar.button(f'结果{random.random()}:{filename}'):
                text_input = st.text_area(filename, height=350, value=content)
                st.write("hello,world")
                if st.button("打开文件"):
                    subprocess.run(['start', f'{document_path}\\{filename}'], shell=True)
    

