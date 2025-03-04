from langchain.embeddings import HuggingFaceBgeEmbeddings
import os
import shutil
import docx
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain.vectorstores import FAISS
from langchain.schema import Document
import uuid
import time

model_path = "..\\model\\bge-small-zh-v1.5"
documents_path = "..\\resources\\documents"
filelst_path = "..\\resources\\filelst.txt"
vdb_path = "..\\resources\\VDB.db"
data_path = "..\\resources\\data"

def myinit(documents_path,filelst_path):
    if not os.path.exists(documents_path):
        os.makedirs(documents_path)
    if not os.path.exists(data_path):
        os.makedirs(data_path)
    if not os.path.exists(filelst_path):
        with open(filelst_path, "w") as f:
            f.write("")

def read_filenames(file_path):
    with open(file_path, 'r') as file:
        filenames = [line.strip() for line in file]
    return filenames

def find_unique(list1, list2):
    set1 = set(list1)
    set2 = set(list2)
    
    unique_in_list = set1 - set2
    
    return list(unique_in_list)

def split_docx(filepath, chunk_size=256, chunk_overlap=0, separators=[".","。","\n","\n\n","?","!"]):
    doc = docx.Document(filepath)
    raw_document = "\n".join([para.text for para in doc.paragraphs])
    splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap,separators=separators)
    return splitter.split_text(raw_document)

def write_filenames(lst, file_name):
    with open(file_name, 'w') as file:
        for item in lst:
            file.write(str(item) + '\n')


class Myembedding_func:
    def __init__(self,model):
        self.model = model

    def embed_documents(self, paragraphs):
        return self.model.embed_documents(paragraphs)
    
    def __call__(self, paragraph):
        return self.model.embed_query(paragraph)

class MyFaiss:
    def __init__(self,vdb_path,documents_path,filelst_path,embedding_func):
        self.vdb_path = vdb_path
        self.documents_path = documents_path
        self.filelst_path = filelst_path
        self.embedding_func = embedding_func
        self.vector_store = None
        if not os.path.exists(vdb_path):
            doc_id = str(uuid.uuid1())
            # doc_id = uuid.uuid1()
            self.vector_store = FAISS.from_documents([Document(page_content="*", metadata={"file": "test","id":doc_id})],self.embedding_func,ids=[doc_id])
        else:
            self.vector_store = FAISS.load_local(vdb_path,self.embedding_func,allow_dangerous_deserialization=True)
        self.vector_store.save_local(vdb_path)

    def add(self,files2add):    
        for file in files2add:
            print('add',file)
            if not file.endswith(".docx"):
                continue
            texts = split_docx(f'{self.documents_path}\\{file}',chunk_size=128)
            documents = []
            ids = []
            for text in texts:
                doc_id = str(uuid.uuid1())
                # doc_id = uuid.uuid1()
                ids.append(doc_id)
                doc = Document(page_content=text, metadata={"file": file,"id":doc_id})
                documents.append(doc)
            if len(documents) > 0:
                self.vector_store.add_documents(documents,ids=ids)
                self.vector_store.save_local(self.vdb_path)
                write_filenames(ids, f"{data_path}\\{file}.txt")

    def delete(self,files2del):
        for file in files2del:
            print('delete',file)
            id2delete = read_filenames(f"{data_path}\\{file}.txt")
            self.vector_store.delete(id2delete)
            self.vector_store.save_local(self.vdb_path)
            os.remove(f"{data_path}\\{file}.txt")
            # while True:
            #     para2delete = self.vector_store.search("*", k=20,search_type="similarity",filter={"file": file})
            #     # para2delete = self.loop_search(300,filter={"file": file})
            #     id2delete = [para.metadata['id'] for para in para2delete]
            #     #id2delete = self.loop_search(300,filter={"file": file})
            #     if len(id2delete) == 0:
            #         break
            #     self.vector_store.delete(id2delete)
            #     self.vector_store.save_local(self.vdb_path)
            #     time.sleep(0.5)

    def update(self):
        pre_list = read_filenames(self.filelst_path)
        current_list = os.listdir(self.documents_path)

        files2add = find_unique(current_list, pre_list)
        files2del = find_unique(pre_list, current_list)

        self.add(files2add)
        self.delete(files2del)

        write_filenames(current_list, self.filelst_path)
        self.vector_store.save_local(self.vdb_path)

    def search(self,query,topk):
        return self.vector_store.search(query, k=topk,search_type="similarity")
    
    def loop_search(self,n=10,filter=None):
        res = []
        for i in range(n):
            paras = self.vector_store.search("*", k=10,search_type="similarity",filter=filter)
            ids = [para.metadata['id'] for para in paras]
            res.extend(ids)
        return list(set(res))


if __name__ == '__main__':
    # 初始化
    myinit(documents_path,filelst_path)
    #加载model
    model = HuggingFaceBgeEmbeddings(model_name = model_path)
    #加载FAISS
    vdb = MyFaiss(vdb_path,documents_path,filelst_path,Myembedding_func(model))

    vdb.update()
    res = vdb.search("如何制造离心机",10)
    for doc in res:
        print("文档:",doc.metadata["file"])
        # print("内容:",doc.page_content)
        # print("ID:",doc.metadata["id"])
        # print("-------------------------------------------------")
        # print("-------------------------------------------------")
        

    